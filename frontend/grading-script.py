import os
import json
import fitz  # PyMuPDF
import pandas as pd
import ollama
from docx import Document

# ==========================================
# CONFIGURATION
# ==========================================
# Get the directory where this script is located
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

INPUT_FOLDER = os.path.join(SCRIPT_DIR, "student_papers")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "grading_results.xlsx")
THAI_MODEL = "scb10x/typhoon2.1-gemma3-12b"
ASSIGNMENT_FILE = os.path.join(SCRIPT_DIR, "assignment.docx")
RUBRIC_FILE = os.path.join(SCRIPT_DIR, "rubric_thai.docx")


# EXTRACT TEXT FROM DOCX
def extract_text_from_docx(docx_path):
    """
    Opens a .docx file and extracts all text from it, including tables.
    """
    try:
        if not os.path.exists(docx_path):
            print(f"[ERROR] File not found: {docx_path}")
            return ""

        doc = Document(docx_path)
        full_text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"

        # Extract text from tables with better structure
        for table in doc.tables:
            full_text += "\n[TABLE START]\n"
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                full_text += f"Row {row_idx}: {row_text}\n"
            full_text += "[TABLE END]\n"

        return full_text.strip()
    except Exception as e:
        print(f"[ERROR] Error reading DOCX {docx_path}: {e}")
        return ""

# EXTRACT TEXT FROM PDF
def extract_text_from_pdf(pdf_path):
    """
    Opens a PDF and extracts text from all pages.
    """
    try:
        doc = fitz.open(pdf_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text() + "\n"
        return full_text
    except Exception as e:
        print(f"[ERROR] Error reading PDF {pdf_path}: {e}")
        return ""

# EXTRACT DIMENSION NAMES FROM RUBRIC
def extract_dimension_names(rubric_text):
    """
    Extracts the 4 dimension names from the rubric table.
    Returns a list of dimension names.
    """
    import re

    # Multiple patterns to try for extracting dimension names
    patterns = [
        r'^\d+\.\s+([^\(\|]+)(?:\s*\([^)]+\))?',  # Pattern: "1. Name (30%)"
        r'Row\s+\d+:\s*\d+\.\s+([^\(\|]+)(?:\s*\([^)]+\))?',  # Pattern from table: "Row X: 1. Name (30%)"
        r'\|\s*\d+\.\s+([^\(\|]+)(?:\s*\([^)]+\))?',  # Pattern with pipe: "| 1. Name (30%)"
    ]

    dimension_names = []

    for line in rubric_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        for pattern in patterns:
            match = re.search(pattern, line)
            if match:
                dimension_name = match.group(1).strip()
                # Filter out empty names and duplicates
                if dimension_name and len(dimension_name) > 2 and dimension_name not in dimension_names:
                    dimension_names.append(dimension_name)
                    if len(dimension_names) == 4:  # Stop after finding 4 dimensions
                        return dimension_names
                break

    return dimension_names

# AI GRADER (LOCAL TYPHOON)
def grade_student_work(student_text, assignment_details, rubric, dimension_names):
    """
    Sends the text to Ollama (Typhoon 2) and expects JSON back.
    """
    if not student_text.strip():
        return None

    # Build dimension names for the prompt
    dim_1 = dimension_names[0] if len(dimension_names) > 0 else "มิติที่ 1"
    dim_2 = dimension_names[1] if len(dimension_names) > 1 else "มิติที่ 2"
    dim_3 = dimension_names[2] if len(dimension_names) > 2 else "มิติที่ 3"
    dim_4 = dimension_names[3] if len(dimension_names) > 3 else "มิติที่ 4"

    prompt = f"""
    คุณเป็นผู้ช่วยอาจารย์ (Teaching Assistant) หน้าที่ของคุณคือตรวจการบ้าน

    รายละเอียดของการบ้าน (ASSIGNMENT):
    {assignment_details}

    เกณฑ์การให้คะแนน (RUBRIC):
    {rubric}

    เนื้อหาของนักศึกษา (STUDENT WORK):
    {student_text}

    คำสั่งสำคัญ - อ่านให้ละเอียด:

    โครงสร้าง RUBRIC:
    - RUBRIC มี 4 มิติการประเมิน:
      1. {dim_1}
      2. {dim_2}
      3. {dim_3}
      4. {dim_4}
    - แต่ละมิติมี 4 ระดับคุณภาพ: ดีมาก (4), ดี (3), พอใช้ (2), ปรับปรุง (1)
    - คะแนนเต็มของแต่ละมิติคือ 4 คะแนน (ไม่ใช่ค่าเปอร์เซ็นต์ที่ระบุ)
    - ตัวเลขเปอร์เซ็นต์ (30%, 20%) เป็นน้ำหนักสำหรับการคำนวณคะแนนรวมถ่วงน้ำหนัก ไม่ใช่คะแนนเต็ม

    วิธีการให้คะแนน:
    1. อ่านเนื้อหางานของนักศึกษา
    2. สำหรับแต่ละมิติ ให้คะแนนระหว่าง 1-4 เท่านั้น:
       - 4 = ดีมาก (ตรงกับคำอธิบายในคอลัมน์ "ดีมาก")
       - 3 = ดี (ตรงกับคำอธิบายในคอลัมน์ "ดี")
       - 2 = พอใช้ (ตรงกับคำอธิบายในคอลัมน์ "พอใช้")
       - 1 = ปรับปรุง (ตรงกับคำอธิบายในคอลัมน์ "ปรับปรุง")
    3. ห้ามให้คะแนนเกิน 4 หรือน้อยกว่า 1 ในแต่ละมิติ
    4. ให้คะแนนเป็นทศนิยมได้ (เช่น 2.5, 3.5) แต่ต้องอยู่ในช่วง 1.0-4.0 เท่านั้น
    5. คะแนนรวม (total_score) = dimension_1_score + dimension_2_score + dimension_3_score + dimension_4_score
    6. คะแนนรวมต้องอยู่ระหว่าง 4-16 เท่านั้น (4 มิติ × คะแนนเต็มมิติละ 4)

    วิเคราะห์และตอบกลับเป็น JSON เท่านั้น (ห้ามมีคำเกริ่นนำ):
    คุณต้องใช้ชื่อมิติที่ระบุด้านล่างนี้เท่านั้น ห้ามเปลี่ยนแปลงชื่อมิติ
    {{
        "dimension_1_name": "{dim_1}",
        "dimension_1_max": 4,
        "dimension_1_score": คะแนนที่ได้ (1.0-4.0),
        "dimension_2_name": "{dim_2}",
        "dimension_2_max": 4,
        "dimension_2_score": คะแนนที่ได้ (1.0-4.0),
        "dimension_3_name": "{dim_3}",
        "dimension_3_max": 4,
        "dimension_3_score": คะแนนที่ได้ (1.0-4.0),
        "dimension_4_name": "{dim_4}",
        "dimension_4_max": 4,
        "dimension_4_score": คะแนนที่ได้ (1.0-4.0),
        "total_score": ผลรวมของคะแนนทั้ง 4 มิติ (4.0-16.0),
        "strengths": "จุดแข็ง 2-3 ประโยค อธิบายให้ละเอียด",
        "weaknesses": "จุดอ่อน 2-3 ประโยค อธิบายให้ละเอียด",
        "feedback": "คำแนะนำเพิ่มเติม 3-5 ประโยค อธิบายอย่างละเอียดว่านักศึกษาควรปรับปรุงอย่างไร"
    }}
    """

    try:
        response = ollama.chat(model=THAI_MODEL, messages=[
            {'role': 'user', 'content': prompt}
        ], format='json')

        return json.loads(response['message']['content'])
    except Exception as e:
        print(f"[ERROR] AI Error: {e}")
        return None


# MAIN EXECUTION LOOP
def main():
    # Load assignment details and rubric from .docx files
    print("Loading assignment details and rubric...")
    assignment_details = extract_text_from_docx(ASSIGNMENT_FILE)
    rubric = extract_text_from_docx(RUBRIC_FILE)

    if not assignment_details:
        print(f"[ERROR] Failed to load assignment from '{ASSIGNMENT_FILE}'. Exiting.")
        return

    if not rubric:
        print(f"[ERROR] Failed to load rubric from '{RUBRIC_FILE}'. Exiting.")
        return

    print("[OK] Assignment and rubric loaded successfully.\n")

    # Extract dimension names from rubric
    dimension_names = extract_dimension_names(rubric)
    if len(dimension_names) < 4:
        print(f"[WARNING] Could not extract all 4 dimension names. Found {len(dimension_names)} dimensions.")
        # Use default names if extraction fails
        dimension_names = ["มิติที่ 1", "มิติที่ 2", "มิติที่ 3", "มิติที่ 4"]
    print(f"[OK] Extracted {len(dimension_names)} dimension names from rubric.\n")

    results = []

    # Get all PDF files in the folder
    files = [f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith('.pdf')]
    print(f"Found {len(files)} papers in '{INPUT_FOLDER}'...")

    for filename in files:
        print(f"   Processing: {filename} ...", end=" ", flush=True)

        file_path = os.path.join(INPUT_FOLDER, filename)

        # 1. Extract Text
        text = extract_text_from_pdf(file_path)

        if len(text) < 50:
            print("[SKIP] (Empty or Image-based PDF)")
            continue

        # 2. Grade with AI
        grading = grade_student_work(text, assignment_details, rubric, dimension_names)

        if grading:
            print("[DONE]")
            results.append({
                "Student_File": filename,
                "Dimension_1": grading.get('dimension_1_name', '-'),
                "Max_Score_1": grading.get('dimension_1_max', 0),
                "Score_1": grading.get('dimension_1_score', 0),
                "Dimension_2": grading.get('dimension_2_name', '-'),
                "Max_Score_2": grading.get('dimension_2_max', 0),
                "Score_2": grading.get('dimension_2_score', 0),
                "Dimension_3": grading.get('dimension_3_name', '-'),
                "Max_Score_3": grading.get('dimension_3_max', 0),
                "Score_3": grading.get('dimension_3_score', 0),
                "Dimension_4": grading.get('dimension_4_name', '-'),
                "Max_Score_4": grading.get('dimension_4_max', 0),
                "Score_4": grading.get('dimension_4_score', 0),
                "Total_Score": grading.get('total_score', 0),
                "Strengths": grading.get('strengths', '-'),
                "Weaknesses": grading.get('weaknesses', '-'),
                "Feedback": grading.get('feedback', '-')
            })
        else:
            print("[FAILED]")

    # 3. Save to Excel
    if results:
        df = pd.DataFrame(results)
        df.to_excel(OUTPUT_FILE, index=False)
        print(f"\n[SUCCESS] Results saved to '{OUTPUT_FILE}'")
    else:
        print("\n[WARNING] No papers were successfully graded.")

if __name__ == "__main__":
    main()